import streamlit as st
import pandas as pd
from supabase import create_client, Client
import os
import time
import json
import io
import textwrap
from datetime import datetime, timezone, timedelta
from dotenv import load_dotenv
from types import SimpleNamespace
import extra_streamlit_components as xtc
try:
    from streamlit_autorefresh import st_autorefresh
except ImportError:
    st_autorefresh = None

# Importación segura de librerías opcionales
try:
    import openpyxl
    HAS_OPENPYXL = True
except (ImportError, ModuleNotFoundError):
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Helper para zona horaria (Lima/Bogotá UTC-5)
def get_lima_now():
    return datetime.now(timezone.utc) - timedelta(hours=5)

def generate_word_letter(texto_completo, firma_resp):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Simular membrete simple
    header = doc.sections[0].header
    htable = header.add_table(1, 2, width=Inches(6))
    htable.autofit = False
    htable.columns[0].width = Inches(3)
    htable.columns[1].width = Inches(3)
    
    # Contenido del cuerpo (dividir por saltos de línea para prrafos)
    for paragraph in texto_completo.split('\n'):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Firma
    doc.add_paragraph("\n\n")
    p_firma = doc.add_paragraph(firma_resp)
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cargo = doc.add_paragraph("Responsable")
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Cargar variables del archivo .env buscando el archivo en la misma carpeta que este script
env_path = os.path.join(os.path.dirname(__file__), '.env')
if os.path.exists(env_path):
    load_dotenv(env_path)
else:
    load_dotenv() # Fallback por si acaso

# Configuración de la página
st.set_page_config(
    page_title="Control Horas - ER",
    page_icon="⌚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Inicializacin de Supabase con soporte para Nube
@st.cache_resource
def get_supabase():
    # 1. Intentar cargar desde Secrets o entorno
    url = st.secrets.get("SUPABASE_URL") or os.getenv("SUPABASE_URL")
    key = st.secrets.get("SUPABASE_KEY") or os.getenv("SUPABASE_KEY")
    service_key = st.secrets.get("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_SERVICE_KEY")

    if not url or not key:
        st.error(" Configuracin incompleta. Revisa los Secrets de Streamlit.")
        st.stop()

    # 2. Limpieza de llaves
    def clean(v):
        return str(v).strip().strip('"').strip("'").strip() if v else None

    url, key, service_key = map(clean, [url, key, service_key])
    
    # Priorizar Service Key para administración
    return create_client(url, service_key if service_key else key)

supabase = get_supabase()
# Inicializar gestor de cookies (CRITICAL PARA IOS)
cookie_manager = xtc.CookieManager()

# --- TRAMPA DE CIERRE TOTAL (HARD LOGOUT) ---
# Si detectamos 'logout' en la URL, limpiamos TODO antes de que la app se cargue
if st.query_params.get("logout") == "1":
    st.info("🔄 Limpiando sesión de forma segura...")
    # Limpiar Python
    st.session_state.user = None
    st.session_state.logout_requested = True
    st.components.v1.html("""
        <script>
            localStorage.clear();
            document.cookie.split(";").forEach(function(c) {
                // Borrado universal de cookies por path y dominio
                const base = c.replace(/^ +/, "").replace(/=.*/, "=;expires=" + new Date().toUTCString() + ";path=/");
                document.cookie = base;
            });
        </script>
    """, height=0)
    time.sleep(1.0)
    st.query_params.clear() # Limpiar el ?logout=1
    st.rerun() # Esto nos llevará a la pantalla de login limpia

# Estilos premium
st.markdown("""
<style>
    .main { background-color: #f5f3ff; }
    .stButton>button { background-color: #7c3aed; color: white; border-radius: 8px; font-weight: bold; width: 100%; }
    .stButton>button:hover { background-color: #6d28d9; border-color: #6d28d9; }
    .sidebar .sidebar-content { background-color: #ede9fe; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] { 
        background-color: #f3f4f6; padding: 10px 20px; border-radius: 8px 8px 0 0;
    }
    .stTabs [aria-selected="true"] { background-color: #7c3aed !important; color: white !important; }
</style>
""", unsafe_allow_html=True)

# Lógica de Login
def login_user(email, password):
    try:
        response = supabase.auth.sign_in_with_password({"email": email, "password": password})
        if response.user:
            # Fetch profile with roles
            try:
                profile = supabase.table("profiles").select("*, roles(name)").eq("id", response.user.id).single().execute()
                p_data = profile.data if profile else None
            except Exception as pe:
                st.error(f" Error recuperando perfil: {str(pe)}")
                return
            
            if not p_data or not p_data.get('is_active', False):
                st.error(" Usuario inexistente o desactivado. Contacte al administrador.")
                return

            st.session_state.user = response.user
            st.session_state.profile = p_data
            is_admin_check = p_data.get('is_admin', False)
            acc_type = p_data.get('account_type', '')
            st.session_state.is_admin = is_admin_check or (acc_type == "Administrador")
            st.session_state.logout_requested = False
            
            # Persistencia Universal: document.cookie (Lax) + LocalStorage (Fallback fuerte)
            st.components.v1.html(f"""
                <script>
                    const expire = new Date();
                    expire.setTime(expire.getTime() + (30*24*60*60*1000));
                    const value = "user_id_persist={response.user.id}; expires=" + expire.toUTCString() + "; path=/; SameSite=Lax";
                    document.cookie = value;
                    localStorage.setItem("user_id_persist", "{response.user.id}");
                </script>
            """, height=0)
            
            
            st.success(" Sesión guardada en este dispositivo.")
            time.sleep(1.0)
            # Limpiar cualquier estado residual de cronómetros anteriores
            limpiar_estado_timer()
            st.rerun()
    except Exception as e:
        st.error(f" Error de acceso: {str(e)}")

def check_overlap(user_id, start_dt, end_dt):
    """Validar que no existan registros superpuestos para el mismo usuario.
    Dos rangos se solapan si: (start1 < end2) AND (end1 > start2)
    """
    try:
        # Convertir a UTC y luego a strings ISO SIN timezone
        if hasattr(start_dt, 'tzinfo'):
            if start_dt.tzinfo is None:
                start_utc = start_dt.replace(tzinfo=timezone.utc)
                end_utc = end_dt.replace(tzinfo=timezone.utc)
            else:
                start_utc = start_dt.astimezone(timezone.utc)
                end_utc = end_dt.astimezone(timezone.utc)
        else:
            start_utc = start_dt.replace(tzinfo=timezone.utc)
            end_utc = end_dt.replace(tzinfo=timezone.utc)
        
        # Convertir a ISO strings SIN timezone para comparación
        start_iso = start_utc.replace(tzinfo=None).isoformat()
        end_iso = end_utc.replace(tzinfo=None).isoformat()
        
        # Crear timestamp de inicio del día (00:00:00) como string ISO
        day_start = start_utc.replace(hour=0, minute=0, second=0, microsecond=0, tzinfo=None).isoformat()
        
        # Obtener todos los registros del usuario desde el inicio del día
        q = supabase.table("time_entries").select("id, start_time, end_time").eq("profile_id", user_id).gte("start_time", day_start).execute()
        
        if not q.data:
            return False
        
        # Verificar solapamiento comparando strings ISO
        for entry in q.data:
            existing_start = entry['start_time']
            existing_end = entry['end_time']
            
            # Lógica de solapamiento con strings ISO
            if (existing_start < end_iso) and (existing_end > start_iso):
                return True
        
        return False
    except Exception as e:
        # En caso de error, permitir el registro (fail-safe)
        return False

@st.cache_data(ttl=300, show_spinner=False)
def get_clientes_cached():
    try:
        return supabase.table("clients").select("id, name").order("name").execute()
    except:
        return None

# Sidebar y Ttulo
st.title(" Control Horas - ER")

if 'user' not in st.session_state:
    st.session_state.user = None
if 'logout_requested' not in st.session_state:
    st.session_state.logout_requested = False

# Función reutilizable para el Registro de Tiempos
def mostrar_registro_tiempos():
    # Manejo de mensajes persistentes tras rerun
    if 'success_msg' in st.session_state:
        st.toast(st.session_state.success_msg, icon="✅")
        del st.session_state.success_msg
    
    # Manejo de keys para borrado
    if 'form_key_suffix' not in st.session_state: st.session_state.form_key_suffix = 0
    if 'timer_running' not in st.session_state: st.session_state.timer_running = False
    if 'timer_start' not in st.session_state: st.session_state.timer_start = None
    if 'total_elapsed' not in st.session_state: st.session_state.total_elapsed = 0
    if 'active_timer_id' not in st.session_state: st.session_state.active_timer_id = None

    # --- SINCRONIZACIN INICIAL (CRITICAL PARA IOS) ---
    # Se hace AQU para que cargue Cliente/Proyecto ANTES de renderizar el formulario
    if st.session_state.active_timer_id is None and st.session_state.user:
        try:
            timer_q = supabase.table("active_timers").select("*, projects(name, client_id, clients(name))").eq("user_id", st.session_state.user.id).execute()
            if timer_q and timer_q.data:
                t_data = timer_q.data[0]
                
                # --- HEARTBEAT & AUTO-STOP CHECK ---
                # Verificar si el cronómetro está "vivo" o si murió (batería, cierre inesperado)
                should_auto_stop = False
                last_update = pd.to_datetime(t_data.get('updated_at', t_data['created_at'])).replace(tzinfo=timezone.utc)
                now_utc = datetime.now(timezone.utc)
                
                # Si pasaron más de 5 minutos desde último update, asumimos muerte súbita
                if t_data['is_running'] and (now_utc - last_update).total_seconds() > 300:
                    should_auto_stop = True
                    # Calcular tiempo real hasta el corte
                    start_utc = pd.to_datetime(t_data['start_time']).replace(tzinfo=timezone.utc)
                    # Tiempo corrido hasta el último latido
                    valid_elapsed = t_data['total_elapsed_seconds'] + (last_update - start_utc).total_seconds()
                    
                    try:
                        supabase.table("active_timers").update({
                            "is_running": False,
                            "total_elapsed_seconds": int(valid_elapsed),
                            "updated_at": now_utc.isoformat()
                        }).eq("id", t_data['id']).execute()
                        st.toast(f"⚠️ Cronómetro detenido automáticamente (Inactividad desde {last_update.astimezone(timezone(timedelta(hours=-5))).strftime('%H:%M')})", icon="🛑")
                        # Actualizar estado local
                        t_data['is_running'] = False
                        t_data['total_elapsed_seconds'] = int(valid_elapsed)
                    except Exception as e:
                        st.error(f"Error auto-deteniendo cronómetro: {e}")

                # Cargar en sesión
                if st.session_state.active_timer_id != t_data['id'] or should_auto_stop:
                    st.session_state.active_timer_id = t_data['id']
                    st.session_state.active_project_id = t_data['project_id']
                    st.session_state.timer_running = t_data['is_running']
                    st.session_state.active_timer_description = t_data.get('description', '')
                    st.session_state.active_timer_billable = t_data.get('is_billable', True)
                    st.session_state.total_elapsed = t_data['total_elapsed_seconds']
                    st.session_state.timer_start = pd.to_datetime(t_data['start_time']).replace(tzinfo=None) # Local time logic used elsewhere expects naive or handle with care

            else:
                # No active timer found in DB
                st.session_state.active_timer_id = None
                st.session_state.timer_running = False
                st.session_state.active_timer_description = ""
            st.rerun()
        except: pass
    
    # 1. Selección de Cliente (Siempre visible)
    clientes_resp = None
    for _ in range(3): # Simple retry logic
        clientes_resp = get_clientes_cached()
        if clientes_resp: break
        time.sleep(0.5)
        
    if not clientes_resp or not clientes_resp.data:
        st.info("Aún no hay clientes registrados (o error de conexión).")
        return
        
    client_map = {c['name']: c['id'] for c in clientes_resp.data}
    
    # Recuperar cliente del timer activo si existe
    index_cliente = 0
    if 'active_client_name' in st.session_state and st.session_state.active_client_name in client_map:
        index_cliente = (list(client_map.keys()).index(st.session_state.active_client_name)) + 1

    cliente_sel = st.selectbox("Seleccionar Cliente", ["---"] + list(client_map.keys()), index=index_cliente, key=f"cli_{st.session_state.form_key_suffix}")
    
    # Variables de control y estado (Scope de la función)
    fecha_sel = get_lima_now()
    listo_para_registro = False
    p_id = None
    moneda = "$"
    can_register = False
    form_valido = False
    descripcion = ""
    es_facturable = True
    target_user_id = st.session_state.user.id

    if cliente_sel != "---":
        proyectos = supabase.table("projects").select("id, name, currency").eq("client_id", client_map[cliente_sel]).order("name").execute()
        if not proyectos.data:
            st.warning(f"Sin proyectos para {cliente_sel}.")
        else:
            proj_map = {p['name']: p['id'] for p in proyectos.data}
            proj_currency = {p['id']: p['currency'] for p in proyectos.data}
            
            index_proj = 0
            if 'active_project_name' in st.session_state and st.session_state.active_project_name in proj_map:
                index_proj = list(proj_map.keys()).index(st.session_state.active_project_name)

            proyecto_sel = st.selectbox("Seleccionar Proyecto", list(proj_map.keys()), index=index_proj, key=f"pro_{st.session_state.form_key_suffix}")
            p_id = proj_map[proyecto_sel]
            moneda = proj_currency[p_id]
            listo_para_registro = True
            
            st.info(f"Proyecto: **{proyecto_sel}** | Moneda: **{moneda}**")
            
            col_u1, col_u2 = st.columns(2)
            with col_u1:
                fecha_sel = st.date_input("Fecha", value=get_lima_now(), max_value=get_lima_now(), key=f"fec_{st.session_state.form_key_suffix}")
            with col_u2:
                if st.session_state.is_admin:
                    usuarios_res = supabase.table("profiles").select("id, full_name").eq("is_active", True).execute()
                    user_map = {u['full_name']: u['id'] for u in usuarios_res.data}
                    usuario_para = st.selectbox("Registrar para", list(user_map.keys()), index=list(user_map.values()).index(st.session_state.user.id) if st.session_state.user.id in user_map.values() else 0, key=f"user_sel_{st.session_state.form_key_suffix}")
                    target_user_id = user_map[usuario_para]
                else:
                    target_user_id = st.session_state.user.id
                st.write(f"Usuario: **{st.session_state.profile['full_name']}**")

            # VALIDACIÓN DE TARIFA
            if target_user_id:
                try:
                    profile_info = supabase.table("profiles").select("role_id").eq("id", target_user_id).single().execute()
                    role_id = profile_info.data['role_id']
                    rate_q = supabase.table("project_rates").select("rate").eq("project_id", p_id).eq("role_id", role_id).execute()
                    
                    current_rate_val = float(rate_q.data[0]['rate']) if rate_q.data else 0.0
                    if current_rate_val <= 0:
                        if st.session_state.is_admin:
                            st.warning(f" **Atención**: No se han definido tarifas para el rol en este proyecto.")
                    else:
                        if st.session_state.is_admin:
                            st.success(f"Tarifa detectada: **{current_rate_val} {moneda}/h**")
                    can_register = True
                except:
                    pass

            # Valor por defecto para descripción y facturabilidad
            def_desc = st.session_state.get('active_timer_description', '')
            def_fact = st.session_state.get('active_timer_billable', True)

            descripcion = st.text_area("Detalle del trabajo", value=def_desc, placeholder="¿Qué hiciste?", key=f"desc_{st.session_state.form_key_suffix}")
            es_facturable = st.checkbox("Es facturable?", value=def_fact, key=f"fact_{st.session_state.form_key_suffix}")
            form_valido = len(descripcion.strip()) > 3
            if not form_valido:
                st.warning(" ⚠️ Ingrese el **Detalle del trabajo** para habilitar el registro.")

            # Nota Interna (Opcional)
            nota_interna = st.text_input("Nota Interna / Flag (Opcional, solo admins)", key=f"note_{st.session_state.form_key_suffix}")

    st.markdown("---")
    
    # 2. SECCIÓN DE REGISTRO (Solo si hay proyecto)
    if listo_para_registro:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader(" Ingreso Manual")
            t_inicio_str = st.text_input("Hora Inicio (HH:mm)", value="08:00", key=f"hi_{st.session_state.form_key_suffix}")
            t_fin_str = st.text_input("Hora Final (HH:mm)", value="", placeholder="Vacío = Hora Actual (Solo Hoy)", key=f"hf_{st.session_state.form_key_suffix}")
            
            if st.button("Registrar Manualmente", disabled=not (can_register and form_valido), use_container_width=True):
                try:
                    is_today = fecha_sel == get_lima_now().date()
                    tz_local = timezone(timedelta(hours=-5))

                    t1_dt = datetime.strptime(t_inicio_str, "%H:%M")
                    t1 = datetime.combine(fecha_sel, t1_dt.time()).replace(tzinfo=tz_local).astimezone(timezone.utc)
                    
                    if not t_fin_str:
                        if is_today:
                             t2 = get_lima_now() # Ya tiene timezone utc diff correcto si usa .now(timezone.utc) o similar, pero get_lima_now tiene -5
                             # get_lima_now returns now in Lima (-5). We need to ensure t1 is compared correctly.
                             # t1 is converted to UTC. get_lima_now is Lima time.
                             # Let's convert get_lima_now to UTC for storage
                             t2 = t2.astimezone(timezone.utc)
                        else:
                            st.error("Debe especificar Hora Final para fechas pasadas.")
                            t2 = None # Blocking
                    else:
                        t2_dt = datetime.strptime(t_fin_str, "%H:%M")
                        t2 = datetime.combine(fecha_sel, t2_dt.time()).replace(tzinfo=tz_local).astimezone(timezone.utc)
                    
                    if t2:
                        if t2 <= t1:
                            st.error("La hora final debe ser posterior a la inicial.")
                        elif check_overlap(target_user_id, t1, t2):
                            st.error("⚠️ Error: El rango de horas se cruza con un registro existente.")
                        else:
                            supabase.table("time_entries").insert({
                                "profile_id": target_user_id, "project_id": p_id, "description": descripcion,
                                "start_time": t1.isoformat(), "end_time": t2.isoformat(),
                                "total_minutes": int((t2 - t1).total_seconds() / 60), "is_billable": es_facturable,
                                "internal_note": nota_interna
                            }).execute()
                            limpiar_estado_timer()
                            st.session_state.success_msg = f" Guardado con éxito ({t_inicio_str} a {t2.astimezone(tz_local).strftime('%H:%M')})."
                            st.rerun()
                except ValueError:
                    st.error("Formato invlido. Use HH:mm (ej: 08:33)")
                except Exception as e:
                    st.error(f" Error: {str(e)}")

        with col2:
            st.subheader(" Cronómetro")
            is_today = fecha_sel == get_lima_now().date()
            if not is_today and not st.session_state.timer_running:
                 st.info("⚠️ El cronómetro solo está disponible para registros del día de hoy.")
            else:
                timer_is_for_current_proj = (st.session_state.active_timer_id and st.session_state.get('active_project_id') == p_id)

                if st.session_state.timer_running and timer_is_for_current_proj:
                    if st_autorefresh:
                        count = st_autorefresh(interval=50 * 1000, key="timer_pulse")
                        try:
                            now_utc = datetime.now(timezone.utc)
                            supabase.table("active_timers").update({"updated_at": now_utc.isoformat()}).eq("id", st.session_state.active_timer_id).execute()
                        except: pass
                    # -----------------------

                    now_lima = get_lima_now().replace(tzinfo=None)
                    actual_elapsed = st.session_state.total_elapsed + (now_lima - st.session_state.timer_start).total_seconds()
                    hrs, rem = divmod(int(actual_elapsed), 3600)
                    mins, secs = divmod(rem, 60)
                    st.metric(" EN VIVO", f"{hrs:02d}:{mins:02d}:{secs:02d}")

                    c_t1, c_t2, c_t3 = st.columns(3)
                    with c_t1:
                        if st.button(" || Pausar", use_container_width=True):
                            try:
                                t_now = get_lima_now().replace(tzinfo=None)
                                new_elapsed = st.session_state.total_elapsed + (t_now - st.session_state.timer_start).total_seconds()
                                st.session_state.total_elapsed = new_elapsed
                                st.session_state.timer_running = False
                                supabase.table("active_timers").update({
                                    "is_running": False, "total_elapsed_seconds": int(new_elapsed),
                                    "description": descripcion, "is_billable": es_facturable
                                }).eq("id", st.session_state.active_timer_id).execute()
                                st.rerun()
                            except Exception as e:
                                st.error(f" Error al pausar: {str(e)}")
                    with c_t2:
                        if st.button(" 🔄 Sinc", use_container_width=True, help="Fuerza la actualización si el tiempo se ve estático"): st.rerun()
                    with c_t3:
                        if st.button(" Fin", disabled=not (can_register and form_valido), use_container_width=True, type="primary"):
                            try:
                                t_now = get_lima_now().replace(tzinfo=None)
                                t_sec = st.session_state.total_elapsed + (t_now - st.session_state.timer_start).total_seconds()
                                t_min = int(t_sec // 60) + (1 if t_sec % 60 > 0 else 0)
                                tz_local = timezone(timedelta(hours=-5))
                                t_st_loc = st.session_state.timer_start - timedelta(seconds=st.session_state.total_elapsed)
                                st_dt = datetime.combine(fecha_sel, t_st_loc.time()).replace(tzinfo=tz_local).astimezone(timezone.utc)
                                end_dt = st_dt + timedelta(minutes=t_min)
                                
                                if check_overlap(target_user_id, st_dt, end_dt):
                                    st.error("⚠️ Error: El rango de horas se cruza con un registro existente.")
                                else:
                                    # 1. Intentar GUARDAR el registro
                                    insert_ok = False
                                    try:
                                        payload = {
                                            "profile_id": target_user_id, "project_id": p_id, "description": descripcion,
                                            "start_time": st_dt.isoformat(), "end_time": end_dt.isoformat(),
                                            "total_minutes": t_min, "is_billable": es_facturable
                                        }
                                        if nota_interna:
                                            payload["internal_note"] = nota_interna
                                        
                                        supabase.table("time_entries").insert(payload).execute()
                                        insert_ok = True
                                    except Exception as e:
                                        st.error(f" Error al guardar registro: {str(e)}")

                                    # 2. Si guardó, limpiar cronómetro (con fallback)
                                    if insert_ok:
                                        if st.session_state.active_timer_id:
                                            try:
                                                supabase.table("active_timers").delete().eq("id", st.session_state.active_timer_id).execute()
                                            except Exception as e_del:
                                                # Fallback: Intentar al menos detenerlo para que no siga contando
                                                try:
                                                    supabase.table("active_timers").update({"is_running": False}).eq("id", st.session_state.active_timer_id).execute()
                                                except: pass
                                                st.toast(f"Guardado, pero error limpiando timer: {str(e_del)}", icon="⚠️")
                                        
                                        limpiar_estado_timer()
                                        st.session_state.success_msg = " Cronómetro guardado."
                                        st.rerun()
                            except Exception as e:
                                st.error(f" Error inesperado al finalizar: {str(e)}")
                else:
                    if st.session_state.total_elapsed > 0:
                        hrs, rem = divmod(int(st.session_state.total_elapsed), 3600)
                        mins, secs = divmod(rem, 60)
                        st.metric("Pausado", f"{hrs:02d}:{mins:02d}:{secs:02d}")
                        cp1, cp2 = st.columns(2)
                        with cp1:
                            if st.button(" Continuar"):
                                try:
                                    st.session_state.timer_start = get_lima_now().replace(tzinfo=None)
                                    st.session_state.timer_running = True
                                    supabase.table("active_timers").update({
                                        "is_running": True, "start_time": st.session_state.timer_start.isoformat(),
                                        "description": descripcion, "is_billable": es_facturable
                                    }).eq("id", st.session_state.active_timer_id).execute()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f" Error al continuar: {str(e)}")
                        with cp2:
                            if st.button(" Descartar"):
                                try:
                                    if st.session_state.active_timer_id:
                                        supabase.table("active_timers").delete().eq("id", st.session_state.active_timer_id).execute()
                                    limpiar_estado_timer()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f" Error al descartar: {str(e)}")
                    else:
                        if st.button(" Iniciar Cronómetro", disabled=not (can_register and form_valido and is_today)):
                            if not is_today:
                                st.error("Solo se permite iniciar cronómetro hoy.")
                                st.stop()

                            try:
                                st.session_state.timer_start = get_lima_now().replace(tzinfo=None)
                                st.session_state.timer_running = True
                                resp = supabase.table("active_timers").insert({
                                    "user_id": st.session_state.user.id, "project_id": p_id,
                                    "start_time": st.session_state.timer_start.isoformat(),
                                    "description": descripcion, "is_billable": es_facturable, "is_running": True
                                }).execute()
                                if resp.data:
                                    st.session_state.active_timer_id = resp.data[0]['id']
                                    st.session_state.active_project_id = p_id
                                st.rerun()
                            except Exception as e:
                                # Intento de recuperación si falla por duplicado (RLS/Unique violation)
                                if "violates row-level security" in str(e) or "duplicate key" in str(e) or "42501" in str(e):
                                    try:
                                        # Forzar recuperación
                                        rec_q = supabase.table("active_timers").select("*").eq("user_id", st.session_state.user.id).execute()
                                        if rec_q and rec_q.data:
                                            t_rec = rec_q.data[0]
                                            st.session_state.active_timer_id = t_rec['id']
                                            st.session_state.timer_running = t_rec['is_running']
                                            st.session_state.timer_start = pd.to_datetime(t_rec['start_time']).replace(tzinfo=None)
                                            st.session_state.total_elapsed = t_rec['total_elapsed_seconds']
                                            st.session_state.active_project_id = t_rec['project_id']
                                            st.rerun()
                                    except:
                                        pass
                                st.error(f"Error iniciando cronómetro: {str(e)}")
                                if st.button("🔴 Forzar Reinicio de Estado"):
                                    limpiar_estado_timer()
                                    st.rerun()

    # 4. TABLA DE HISTORIAL (Siempre visible al final)
    st.markdown("---")
    mostrar_historial_tiempos()

def limpiar_estado_timer():
    if 'form_key_suffix' not in st.session_state: st.session_state.form_key_suffix = 0
    st.session_state.form_key_suffix += 1
    st.session_state.timer_running = False
    st.session_state.total_elapsed = 0
    st.session_state.timer_start = None
    st.session_state.active_timer_id = None
    st.session_state.active_project_id = None
    st.session_state.active_project_name = None
    st.session_state.active_client_name = None
    st.session_state.active_timer_description = ''
    st.session_state.active_timer_billable = True

def mostrar_historial_tiempos():
    st.subheader(" Historial de Horas")
    query = supabase.table("time_entries").select("*, profiles(full_name, role_id, roles(name)), projects(name, currency, clients(name))").order("start_time", desc=True)
    if not st.session_state.is_admin:
        # Convertir a UTC y remover timezone para evitar errores de Supabase
        limite_30_dias_dt = get_lima_now() - timedelta(days=30)
        limite_30_dias = limite_30_dias_dt.astimezone(timezone.utc).replace(tzinfo=None).isoformat()
        query = query.eq("profile_id", st.session_state.user.id).gte("start_time", limite_30_dias)
    
    entries_resp = query.execute()
    if entries_resp.data:
        df = pd.json_normalize(entries_resp.data)
        def to_local(s):
            if pd.isna(s) or not s: return None
            try: return pd.to_datetime(s, utc=True).tz_convert('America/Lima').tz_localize(None)
            except: return pd.to_datetime(s).replace(tzinfo=None) - pd.Timedelta(hours=5)

        df['dt_ref'] = df['start_time'].fillna(df['created_at'])
        df['dt_st'] = df['dt_ref'].apply(to_local)
        df['dt_en'] = df['end_time'].apply(to_local)
        df['Inicio'] = df['dt_st'].dt.strftime('%H:%M').fillna('---')
        df['Fin'] = df['dt_en'].dt.strftime('%H:%M').fillna('---')
        df['Fecha'] = df['dt_st'].dt.strftime('%d.%m-%Y').fillna('---')
        df['Tiempo'] = df['total_minutes'].apply(lambda x: f"{int(x)//60:02d}:{int(x)%60:02d}")
        df['Cliente'] = df['projects.clients.name'].fillna('...')
        df['Proyecto'] = df['projects.name'].fillna('...')
        df['Moneda'] = df['projects.currency'].fillna('')
        df['Cliente'] = df['projects.clients.name'].fillna('...')
        df['Proyecto'] = df['projects.name'].fillna('...')
        df['Moneda'] = df['projects.currency'].fillna('')
        df['Usuario'] = df['profiles.full_name'].fillna('...')
        # Robust against missing column
        if 'internal_note' in df.columns:
            df['Nota'] = df['internal_note'].fillna('')
        else:
            df['Nota'] = ''
        
        if st.session_state.is_admin:
            rates_resp = supabase.table("project_rates").select("*").execute()
            rates_df = pd.DataFrame(rates_resp.data)
            def calc_metrics(row):
                rate = 0.0
                if not rates_df.empty:
                    r = rates_df[(rates_df['project_id'] == row['project_id']) & (rates_df['role_id'] == row['profiles.role_id'])]
                    rate = float(r['rate'].iloc[0]) if not r.empty else 0.0
                total = (row['total_minutes'] / 60) * rate
                return pd.Series([rate, total, total if row['is_billable'] else 0.0])
            df[['Costo Hora', 'Total Bruto', 'Monto Facturable']] = df.apply(calc_metrics, axis=1)
            
            display_cols = ['Fecha', 'Usuario', 'Cliente', 'Proyecto', 'Moneda', 'description', 'Nota', 'Inicio', 'Fin', 'Tiempo', 'Costo Hora', 'is_billable', 'Total Bruto', 'Monto Facturable', 'is_paid', 'invoice_number']
            final_df = df[display_cols].rename(columns={'description': 'Detalle', 'is_billable': '¿Fact?', 'is_paid': 'Cobrado?'})
            
            edited_df = st.data_editor(final_df, use_container_width=True, hide_index=True,
                column_config={
                    "¿Fact?": st.column_config.CheckboxColumn(label="Fact?"),
                    "Costo Hora": st.column_config.NumberColumn(format="%.2f"),
                    "Total Bruto": st.column_config.NumberColumn(format="%.2f"),
                    "Monto Facturable": st.column_config.NumberColumn(format="%.2f"),
                    "Nota": st.column_config.TextColumn("Nota Interna")
                },
                disabled=['Fecha', 'Usuario', 'Cliente', 'Proyecto', 'Moneda', 'Detalle', 'Inicio', 'Fin', 'Tiempo', 'Costo Hora', 'Total Bruto', '¿Fact?', 'Monto Facturable'])
            
            if st.button("Guardar Cambios Historial"):
                for idx, row in edited_df.iterrows():
                    orig = df.iloc[idx]
                    if row['Cobrado?'] != orig['is_paid'] or row['invoice_number'] != orig['invoice_number'] or row['Nota'] != orig['internal_note']:
                        supabase.table("time_entries").update({
                            "is_paid": row['Cobrado?'], 
                            "invoice_number": row['invoice_number'],
                            "internal_note": row['Nota']
                        }).eq("id", orig['id']).execute()
                st.success("Cambios guardados.")
                st.rerun()
        else:
            view_cols = ['Fecha', 'Cliente', 'Proyecto', 'description', 'Inicio', 'Fin', 'Tiempo']
            st.dataframe(df[view_cols].rename(columns={'description': 'Detalle'}), use_container_width=True, hide_index=True)
    else:
        st.info("No hay registros recientes.")

# --- RECUERDO DE SESIÓN ---
if not st.session_state.user and not st.session_state.get('logout_requested'):
    # 1. Puerta de hidratación para componentes de almacenamiento
    if "init_gate" not in st.session_state:
        st.session_state.init_gate = True
        with st.spinner("⏳ Conectando..."):
            time.sleep(1.0)
            st.rerun()

    # 2. Intentar recuperar de Cookie de sesión
    u_id = cookie_manager.get('user_id_persist')
    
    if u_id:
        try:
            profile_res = supabase.table("profiles").select("*, roles(name)").eq("id", u_id).single().execute()
            if profile_res and profile_res.data and profile_res.data.get('is_active'):
                st.session_state.user = SimpleNamespace(id=u_id)
                st.session_state.profile = profile_res.data
                st.session_state.is_admin = profile_res.data.get('is_admin', False) or (profile_res.data.get('account_type') == "Administrador")
                # Limpiar estado de timer al restaurar sesión
                limpiar_estado_timer()
                st.rerun()
        except:
            pass

# Eliminar inicialización duplicada
# cookie_manager = xtc.CookieManager()

if not st.session_state.user:
    st.subheader("Acceso al Sistema")
    with st.form("login_form"):
        email = st.text_input("Correo electrónico")
        password = st.text_input("Contraseña", type="password")
        if st.form_submit_button("Entrar"):
            login_user(email, password)
else:
    with st.sidebar:
        st.write(f" **{st.session_state.profile['full_name']}**")
        st.write(f" Rol: {st.session_state.profile['roles']['name']}")
        st.write(f" Tipo: {'Administrador' if st.session_state.is_admin else 'Usuario'}")
        if st.button("🔴 Cerrar Sesión"):
            # Signal de Hard Logout vía URL
            st.query_params["logout"] = "1"
            st.session_state.user = None
            st.session_state.logout_requested = True
            st.rerun()

    if st.session_state.is_admin:
        menu = ["Panel General", "Registro de Tiempos", "Clientes", "Proyectos", "Usuarios", "Roles y Tarifas", "Facturación y Reportes", "Carga Masiva"]
        choice = st.sidebar.selectbox("Seleccione Módulo", menu)

        if choice == "Panel General":
            st.header(" Panel General de Horas")
            
            # Query base (Admin ve todo)
            entries_q = supabase.table("time_entries").select("*, profiles(full_name, role_id, roles(name)), projects(name, currency, clients(name))").order("start_time", desc=True)
            entries = entries_q.execute()
            rates = supabase.table("project_rates").select("*").execute()
            
            if entries.data:
                df = pd.json_normalize(entries.data)
                rates_df = pd.DataFrame(rates.data)
                
                # Conversin horaria manual garantizada (UTC-5)
                df['dt_ref'] = df['start_time'].fillna(df['created_at'])
                df['dt_start'] = df['dt_ref'].apply(lambda x: pd.to_datetime(x, utc=True).tz_convert('America/Lima').tz_localize(None) if pd.notna(x) and x != 'nan' else None)
                df['dt_end'] = df['end_time'].apply(lambda x: pd.to_datetime(x, utc=True).tz_convert('America/Lima').tz_localize(None) if pd.notna(x) and x != 'nan' else None)
                
                df['Hora Inicio'] = df['dt_start'].dt.strftime('%H:%M').fillna('---')
                df['Hora Final'] = df['dt_end'].dt.strftime('%H:%M').fillna('---')
                df['Tiempo (hh:mm)'] = df['total_minutes'].apply(lambda x: f"{int(x)//60:02d}:{int(x)%60:02d}")
                df['Fecha'] = df['dt_start'].dt.strftime('%d.%m-%Y').fillna('---')
                
                # Respaldo si fall el apply (si resultaron nulos pero no deberan)
                if not df.empty and df['Hora Inicio'].iloc[0] == '---' and not df['dt_ref'].isnull().all():
                     df['dt_start'] = (pd.to_datetime(df['dt_ref'], utc=True) - pd.Timedelta(hours=5)).dt.tz_localize(None)
                     df['dt_end'] = (pd.to_datetime(df['end_time'], utc=True) - pd.Timedelta(hours=5)).dt.tz_localize(None)
                     df['Hora Inicio'] = df['dt_start'].dt.strftime('%H:%M').fillna('---')
                     df['Hora Final'] = df['dt_end'].dt.strftime('%H:%M').fillna('---')
                     df['Fecha'] = df['dt_start'].dt.strftime('%d.%m-%Y').fillna('---')
                
                def get_cost(row):
                    if not rates_df.empty:
                        r = rates_df[(rates_df['project_id'] == row['project_id']) & (rates_df['role_id'] == row['profiles.role_id'])]
                        return float(r['rate'].iloc[0]) if not r.empty else 0.0
                    return 0.0

                df['Costo Hora'] = df.apply(get_cost, axis=1)
                df['Valor Total'] = (df['total_minutes'] / 60) * df['Costo Hora']
                df['Costo Facturable'] = df.apply(lambda r: r['Valor Total'] if r['is_billable'] else 0.0, axis=1)
                
                # Renombrar para visualizacin
                df = df.rename(columns={
                    'profiles.full_name': 'Usuario',
                    'profiles.roles.name': 'Rol',
                    'projects.clients.name': 'Cliente',
                    'projects.name': 'Proyecto',
                    'is_billable': 'Facturable'
                })
                
                # Filtros
                col_f1, col_f2 = st.columns(2)
                with col_f1:
                    f_user = st.multiselect("Filtrar por Usuario", df['Usuario'].unique())
                with col_f2:
                    f_client = st.multiselect("Filtrar por Cliente", df['Cliente'].unique())
                
                filtered_df = df.copy()
                if f_user: filtered_df = filtered_df[filtered_df['Usuario'].isin(f_user)]
                if f_client: filtered_df = filtered_df[filtered_df['Cliente'].isin(f_client)]
                
                # Columnas finales (Admin ve todo y puede editar)
                display_cols = ['id', 'Fecha', 'Usuario', 'Rol', 'Cliente', 'Proyecto', 'Hora Inicio', 'Hora Final', 'Tiempo (hh:mm)', 'Costo Hora', 'Valor Total', 'Costo Facturable', 'Facturable']
                
                # FORZAR REDONDEO FSICO EN EL DF PARA EVITAR DECIMALES LARGOS
                filtered_df['Costo Hora'] = filtered_df['Costo Hora'].fillna(0).round(2)
                filtered_df['Valor Total'] = filtered_df['Valor Total'].fillna(0).round(2)
                filtered_df['Costo Facturable'] = filtered_df['Costo Facturable'].fillna(0).round(2)

                # Configuracin de columnas para alineacin y formato
                col_config = {
                    "id": None, # Habilitar ocultamiento real sin error
                    "Costo Hora": st.column_config.NumberColumn(format="%.2f"),
                    "Valor Total": st.column_config.NumberColumn(format="%.2f"),
                    "Costo Facturable": st.column_config.NumberColumn(format="%.2f"),
                    "Facturable": st.column_config.CheckboxColumn(label="")
                }
                
                edited_gen = st.data_editor(
                    filtered_df[display_cols], 
                    column_config=col_config,
                    use_container_width=True, hide_index=True,
                    disabled=['Rol', 'Cliente', 'Proyecto', 'Tiempo (hh:mm)', 'Costo Hora', 'Valor Total', 'Costo Facturable'] # Solo lo bsico y Facturable es editable
                )
                
                # El desmarcado de "Facturable" se refleja en el editor. Recalcular mtricas dinmicas para visualizacin rpida:
                billable_total_live = edited_gen[edited_gen['Facturable'] == True]['Costo Facturable'].sum()
                st.info(f" **Total Facturable Proyectado (en esta vista): {billable_total_live:,.2f}**")
                
                col_btn1, col_btn2 = st.columns([1, 1])
                with col_btn1:
                    if st.button("Guardar cambios en Panel General"):
                        for i, row in edited_gen.iterrows():
                            # Encontrar la fila original por ID
                            orig_id = row['id']
                            orig_row = df[df['id'] == orig_id].iloc[0]
                            
                            updates = {}
                            if row['Facturable'] != orig_row['Facturable']: updates["is_billable"] = row['Facturable']
                            if row['Fecha'] != orig_row['Fecha']: 
                                try:
                                    # Intentar parsear fecha editada
                                    new_d = datetime.strptime(row['Fecha'], '%d.%m-%Y')
                                    # Mantener la hora original
                                    old_dt = pd.to_datetime(orig_row['dt_ref'])
                                    new_dt = old_dt.replace(year=new_d.year, month=new_d.month, day=new_d.day)
                                    updates["start_time"] = new_dt.isoformat()
                                except: pass
                            
                            if updates:
                                supabase.table("time_entries").update(updates).eq("id", orig_id).execute()
                        st.success(" Cambios administrativos guardados.")
                        st.rerun()

                with col_btn2:
                    if HAS_OPENPYXL:
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            filtered_df[display_cols].to_excel(writer, index=False, sheet_name='Historial')
                        st.download_button(
                            label="Descargar Reporte Excel ",
                            data=output.getvalue(),
                            file_name=f"historial_horas_{get_lima_now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    else:
                        st.empty() # No mostrar error si no hay librera
                
                # Calcular inversin por moneda
                st.subheader("Inversin Total por Divisa")
                if not filtered_df.empty:
                    # Agrupar por la moneda del proyecto (que sacamos del join)
                    if 'projects.currency' in filtered_df:
                        metrics_cols = st.columns(len(filtered_df['projects.currency'].unique()))
                        for i, (curr, group) in enumerate(filtered_df.groupby('projects.currency')):
                            with metrics_cols[i]:
                                total_curr = group['Valor Total'].sum()
                                st.metric(f"Total {curr}", f"{curr} {total_curr:,.2f}")
                    else:
                        st.metric("Inversin Total", f"${filtered_df['Valor Total'].sum():,.2f}")
                
                st.markdown("---")
                st.subheader(" Descarga Global de Datos")
                if st.button("Descargar Base de Datos Completa (Excel)"):
                    if HAS_OPENPYXL:
                        try:
                            # Descargar TODO lo que hay en time_entries sin filtros
                            all_q = supabase.table("time_entries").select("*, profiles(full_name), projects(name, currency, clients(name))").execute()
                            if all_q.data:
                                df_all = pd.json_normalize(all_q.data)
                                output_all = io.BytesIO()
                                with pd.ExcelWriter(output_all, engine='openpyxl') as writer:
                                    df_all.to_excel(writer, index=False, sheet_name='BaseCompleta')
                                st.download_button(
                                    label="Confirmar Descarga Global",
                                    data=output_all.getvalue(),
                                    file_name=f"FULL_DB_{get_lima_now().strftime('%Y%m%d_%H%M')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                            else:
                                st.warning("La base de datos est vaca.")
                        except Exception as e:
                            st.error(f"Error en descarga global: {e}")
                    else:
                        st.error("Librera Excel no disponible.")

            else:
                st.info("No hay registros de tiempo an.")

        elif choice == "Registro de Tiempos":
            mostrar_registro_tiempos()

        elif choice == "Clientes":
            st.header(" Gestin de Clientes")
            with st.expander(" Crear Nuevo Cliente", expanded=True):
                with st.form("form_cliente"):
                    nombre = st.text_input("Nombre o Razn Social")
                    doi_type = st.selectbox("Tipo DOI", ["RUC", "DNI", "CE", "PASAPORTE", "OTROS"])
                    doi_num = st.text_input("Nmero de Documento")
                    email_cli = st.text_input("Email de contacto")
                    celular_cli = st.text_input("Nmero de Contacto")
                    direccion = st.text_area("Direccin")
                    
                    if st.form_submit_button("Guardar Cliente"):
                        existente = supabase.table("clients").select("*").or_(f"name.eq.{nombre},doi_number.eq.{doi_num}").execute()
                        if existente.data:
                            st.error(" Error: Ya existe un cliente con ese nombre o número de documento.")
                        else:
                            supabase.table("clients").insert({
                                "name": nombre, "doi_type": doi_type, "doi_number": doi_num, 
                                "address": direccion, "email": email_cli, "contact_number": celular_cli
                            }).execute()
                            st.success(f" Cliente '{nombre}' creado con xito.")
            
            st.subheader("Clientes Registrados")
            clientes_q = supabase.table("clients").select("*").order("name").execute()
            if clientes_q.data:
                c_df = pd.DataFrame(clientes_q.data)
                edited_clients = st.data_editor(
                    c_df[['id', 'name', 'doi_type', 'doi_number', 'email', 'contact_number', 'address']],
                    use_container_width=True, hide_index=True,
                    disabled=["id"]
                )
                if st.button("Guardar Cambios de Clientes"):
                    for i, row in edited_clients.iterrows():
                        orig = clientes_q.data[i]
                        # Solo actualizar si hubo cambios
                        diff = {k: v for k, v in row.items() if v != orig.get(k)}
                        if diff:
                            supabase.table("clients").update(diff).eq("id", row['id']).execute()
                    st.success(" Datos de clientes actualizados.")
                    st.rerun()

        elif choice == "Proyectos":
            st.header(" Gestin de Proyectos")
            clientes = supabase.table("clients").select("id, name").order("name").execute()
            if not clientes.data:
                st.warning("Debe crear un cliente primero.")
            else:
                client_map = {c['name']: c['id'] for c in clientes.data}
                if 'proj_key_suffix' not in st.session_state: st.session_state.proj_key_suffix = 0
                if 'proj_success_msg' in st.session_state:
                    st.success(st.session_state.proj_success_msg)
                    del st.session_state.proj_success_msg
                    
                with st.expander(" Crear Nuevo Proyecto"):
                    with st.form("form_proyecto"):
                        cliente_create = st.selectbox("Seleccionar Cliente", list(client_map.keys()), key=f"p_c_create_{st.session_state.proj_key_suffix}")
                        proj_name = st.text_input("Nombre del Proyecto", key=f"p_name_{st.session_state.proj_key_suffix}")
                        moneda = st.selectbox("Moneda del Proyecto", ["PEN", "USD"], key=f"p_curr_{st.session_state.proj_key_suffix}")
                        if st.form_submit_button("Crear Proyecto"):
                            existente = supabase.table("projects").select("*").eq("client_id", client_map[cliente_create]).eq("name", proj_name).execute()
                            if existente.data:
                                st.error(f" El cliente '{cliente_create}' ya tiene un proyecto llamado '{proj_name}'.")
                            else:
                                supabase.table("projects").insert({
                                    "client_id": client_map[cliente_create], "name": proj_name, "currency": moneda
                                }).execute()
                                st.session_state.proj_success_msg = f" Proyecto '{proj_name}' creado con xito."
                                st.session_state.proj_key_suffix += 1
                                st.rerun()

                st.subheader("Proyectos Existentes")
                proyectos = supabase.table("projects").select("*, clients(name)").order("name").execute()
                if proyectos.data:
                    p_df = pd.json_normalize(proyectos.data)
                    p_df = p_df[['clients.name', 'name', 'currency']]
                    p_df.columns = ['Cliente', 'Proyecto', 'Moneda']
                    st.table(p_df)
                    
                    st.markdown("---")
                    st.subheader(" Editar Moneda de Proyecto")
                    proj_list = {f"{p['clients']['name']} - {p['name']}": p['id'] for p in proyectos.data}
                    p_to_edit = st.selectbox("Seleccionar Proyecto para Editar", list(proj_list.keys()))
                    
                    # Buscar moneda actual
                    curr_p = [p for p in proyectos.data if p['id'] == proj_list[p_to_edit]][0]
                    new_curr = st.selectbox("Nueva Moneda", ["PEN", "USD"], index=0 if curr_p['currency'] == 'PEN' else 1)
                    
                    if st.button("Actualizar Moneda"):
                        supabase.table("projects").update({"currency": new_curr}).eq("id", proj_list[p_to_edit]).execute()
                        st.success(f" Moneda de '{p_to_edit}' actualizada a {new_curr}.")
                        st.rerun()
                else:
                    st.info("No hay proyectos registrados.")

        elif choice == "Usuarios":
            st.header(" Gestin de Usuarios")
            roles = supabase.table("roles").select("id, name").order("name").execute()
            role_map = {r['name']: r['id'] for r in roles.data}
            
            with st.form("form_usuario"):
                u_email = st.text_input("Email (ser su acceso)")
                u_pass = st.text_input("Contraseña", type="password")
                u_name = st.text_input("Nombre Completo")
                u_username = st.text_input("Nombre de Usuario (interno)")
                u_doi_type = st.selectbox("Tipo DOI", ["DNI", "RUC", "CE", "PASAPORTE"])
                u_doi_number = st.text_input("Nmero de DOI")
                u_role = st.selectbox("Rol Operativo (para tarifas)", list(role_map.keys()))
                u_is_admin = st.checkbox("Es Administrador?")
                st.info(" Por seguridad, los nuevos usuarios se crean DESACTIVADOS.")
                
                if st.form_submit_button("Crear Usuario"):
                    if not u_email or not u_pass:
                        st.error(" Email y contraseña son obligatorios.")
                    else:
                        try:
                            new_u = supabase.auth.admin.create_user({
                                "email": u_email.strip(), 
                                "password": u_pass, 
                                "email_confirm": True
                            })
                            supabase.table("profiles").insert({
                                "id": new_u.user.id, 
                                "username": u_username, 
                                "full_name": u_name, 
                                "role_id": role_map[u_role],
                                "doi_type": u_doi_type,
                                "doi_number": u_doi_number,
                                "is_active": False,
                                "is_admin": u_is_admin
                            }).execute()
                            st.success(f" Usuario '{u_name}' creado con xito.")
                            st.info(" Recuerde activarlo en la tabla de abajo para que pueda iniciar sesin.")
                        except Exception as e:
                            st.error(f" Error de permisos: {e}")
                            st.warning("Asegrese de que el 'SUPABASE_SERVICE_KEY' est bien configurado en los Secretos de Streamlit.")
            
            st.subheader("Usuarios Registrados")
            users_resp = supabase.table("profiles").select("*, roles(name)").execute()
            if users_resp.data:
                u_df = pd.DataFrame([
                    {
                        "ID": u['id'],
                        "Nombre": u['full_name'], 
                        "Interno": u['username'],
                        "DOI": f"{u['doi_type']} {u['doi_number']}",
                        "Rol": u['roles']['name'],
                        "Activo": u['is_active'],
                        "Admin": u['is_admin']
                    }
                    for u in users_resp.data
                ])
                
                # Editor para activar/desactivar y dar admin
                st.write("Edite directamente los privilegios en la tabla y guarde:")
                edited_users = st.data_editor(u_df, use_container_width=True, hide_index=True, disabled=["ID", "Interno", "DOI"])
                
                if st.button("Guardar Cambios de Usuarios"):
                    changed_count = 0
                    for i, row in edited_users.iterrows():
                        orig = users_resp.data[i]
                        if row['Activo'] != orig['is_active'] or row['Admin'] != orig['is_admin']:
                            supabase.table("profiles").update({
                                "is_active": row['Activo'], "is_admin": row['Admin']
                            }).eq("id", row['ID']).execute()
                            changed_count += 1
                    if changed_count > 0:
                        st.session_state.user_success_msg = f" {changed_count} usuarios actualizados."
                        st.rerun()

            # Gestin de Mensajes persistentes para usuarios
            if 'user_success_msg' in st.session_state:
                st.success(st.session_state.user_success_msg)
                del st.session_state.user_success_msg

        elif choice == "Roles y Tarifas":
            st.header(" Roles y Tarifas por Proyecto")
            proyectos = supabase.table("projects").select("id, name, clients(name)").execute()
            if not proyectos.data:
                st.warning("Debe crear proyectos primero.")
            else:
                proj_map = {f"{p['clients']['name']} - {p['name']}": p['id'] for p in proyectos.data}
                proj_sel = st.selectbox("Seleccionar Proyecto", list(proj_map.keys()))
                
                # Suffix para resetear inputs de tarifas al cambiar proyecto
                if 'rate_key_prefix' not in st.session_state: st.session_state.rate_key_prefix = 0
                if 'last_proj_sel' not in st.session_state: st.session_state.last_proj_sel = proj_sel
                
                if st.session_state.last_proj_sel != proj_sel:
                    st.session_state.rate_key_prefix += 1
                    st.session_state.last_proj_sel = proj_sel
                    st.rerun()
                
                roles = supabase.table("roles").select("*").execute()
                
                st.subheader(f"Tarifas para: {proj_sel}")
                tarifas_nuevas = {}
                for role in roles.data:
                    col_r1, col_r2 = st.columns([2, 1])
                    with col_r1:
                        st.write(f"Rol: **{role['name']}**")
                    with col_r2:
                        # Buscar tarifa actual - Usamos cache local por proyecto para evitar queries repetitivas si es posible, o simplemente select simple
                        current_rate = supabase.table("project_rates").select("rate").eq("project_id", proj_map[proj_sel]).eq("role_id", role['id']).execute()
                        initial_val = float(current_rate.data[0]['rate']) if current_rate.data else 0.0
                        new_rate = st.number_input(f"Tarifa ({role['name']})", value=initial_val, key=f"rate_{proj_map[proj_sel]}_{role['id']}")
                        tarifas_nuevas[role['id']] = (new_rate, initial_val, bool(current_rate.data))
                
                if st.button("Guardar Todas las Tarifas"):
                    for r_id, (val, old_val, exists) in tarifas_nuevas.items():
                        if val != old_val:
                            if exists:
                                supabase.table("project_rates").update({"rate": val}).eq("project_id", proj_map[proj_sel]).eq("role_id", r_id).execute()
                            else:
                                supabase.table("project_rates").insert({"project_id": proj_map[proj_sel], "role_id": r_id, "rate": val}).execute()
                    st.success(f" Tarifas para '{proj_sel}' guardadas.")
                    st.rerun()
                    st.success(f" Tarifas para '{proj_sel}' guardadas.")
                    st.rerun()

        elif choice == "Carga Masiva":
            st.header(" Carga Masiva de Datos")
            
            # Tabs para diferentes tipos de carga
            upload_tab1, upload_tab2, upload_tab3, upload_tab4 = st.tabs([" Registros de Tiempo", " Clientes", " Proyectos", " Tarifas"])
            
            with upload_tab1:
                st.subheader("Carga Masiva de Registros de Tiempo")
                st.info(" **Formato requerido**: Fecha | Responsable | Cliente | Proyecto | Detalle | Hora Inicio | Hora Final")
                
                # Gestión proactiva de librerías
                if not HAS_OPENPYXL:
                    try:
                        import openpyxl
                        HAS_OPENPYXL = True
                    except:
                        pass
                
                if not HAS_OPENPYXL:
                    st.warning(" ⚠️ La función de carga masiva requiere 'openpyxl'.")
                    if st.button(" 🔄 Re-intentar detectar librerías"):
                        st.rerun()
                
                # Botón para descargar template
                if HAS_OPENPYXL:
                    template_time = pd.DataFrame({
                        'Fecha': ['06.02-2026', '06.02-2026'],
                        'Responsable': ['Juan Pérez', 'Mara García'],
                        'Cliente': ['Cliente A', 'Cliente B'],
                        'Proyecto': ['Proyecto X', 'Proyecto Y'],
                        'Detalle': ['Reunin de planificacin', 'Desarrollo de mdulo'],
                        'Hora Inicio': ['09:00', '14:00'],
                        'Hora Final': ['11:30', '17:00']
                    })
                    buffer_template = io.BytesIO()
                    with pd.ExcelWriter(buffer_template, engine='openpyxl') as writer:
                        template_time.to_excel(writer, index=False, sheet_name='Registros')
                    st.download_button(" Descargar Template", data=buffer_template.getvalue(), file_name="template_registros.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                else:
                    st.warning(" La función de descarga de templates requiere 'openpyxl'. Por favor, instálela.")
                
                uploaded_file = st.file_uploader("Seleccionar archivo Excel", type=['xlsx'], key="upload_time")
                if uploaded_file and HAS_OPENPYXL:
                    try:
                        df_upload = pd.read_excel(uploaded_file)
                        st.write("Vista previa:", df_upload.head())
                        
                        if st.button("Procesar Carga de Registros"):
                            # Mapeos
                            prof_map = {p['full_name']: (p['id'], p['role_id']) for p in supabase.table("profiles").select("id, full_name, role_id").execute().data}
                            clients_map = {c['name']: c['id'] for c in supabase.table("clients").select("id, name").execute().data}
                            
                            success_count = 0
                            errors = []
                            
                            for idx, row in df_upload.iterrows():
                                try:
                                    # Validar usuario
                                    responsable = row.get('Responsable')
                                    if responsable not in prof_map:
                                        errors.append(f"Fila {idx+2}: Responsable '{responsable}' no encontrado")
                                        continue
                                    u_id, role_id = prof_map[responsable]
                                    
                                    # Validar cliente
                                    cliente = row.get('Cliente')
                                    if cliente not in clients_map:
                                        errors.append(f"Fila {idx+2}: Cliente '{cliente}' no encontrado")
                                        continue
                                    c_id = clients_map[cliente]
                                    
                                    # Buscar proyecto
                                    proyecto = row.get('Proyecto')
                                    proj_q = supabase.table("projects").select("id").eq("client_id", c_id).eq("name", proyecto).execute()
                                    if not proj_q.data:
                                        errors.append(f"Fila {idx+2}: Proyecto '{proyecto}' no existe para cliente '{cliente}'")
                                        continue
                                    p_id = proj_q.data[0]['id']
                                    
                                    # Procesar fecha y horas
                                    fecha_str = row.get('Fecha')
                                    if isinstance(fecha_str, datetime):
                                        fecha_dt = fecha_str.date()
                                    else:
                                        fecha_dt = datetime.strptime(str(fecha_str), "%d.%m-%Y").date()
                                    
                                    hora_inicio_str = str(row.get('Hora Inicio'))
                                    hora_final_str = str(row.get('Hora Final'))
                                    
                                    # Parsear horas
                                    t1_dt = datetime.strptime(hora_inicio_str, "%H:%M")
                                    t2_dt = datetime.strptime(hora_final_str, "%H:%M")
                                    
                                    # Crear timestamps UTC-5
                                    tz_local = timezone(timedelta(hours=-5))
                                    t1 = datetime.combine(fecha_dt, t1_dt.time()).replace(tzinfo=tz_local).astimezone(timezone.utc)
                                    t2 = datetime.combine(fecha_dt, t2_dt.time()).replace(tzinfo=tz_local).astimezone(timezone.utc)
                                    
                                    if t2 <= t1:
                                        errors.append(f"Fila {idx+2}: Hora Final debe ser posterior a Hora Inicio")
                                        continue
                                    
                                    # Calcular minutos
                                    total_min = int((t2 - t1).total_seconds() / 60)
                                    
                                    # Insertar
                                    supabase.table("time_entries").insert({
                                        "profile_id": u_id,
                                        "project_id": p_id,
                                        "start_time": t1.isoformat(),
                                        "end_time": t2.isoformat(),
                                        "total_minutes": total_min,
                                        "description": row.get('Detalle', 'Carga Masiva'),
                                        "is_billable": True
                                    }).execute()
                                    success_count += 1
                                    
                                except Exception as e:
                                    errors.append(f"Fila {idx+2}: Error - {str(e)}")
                            
                            st.success(f" Procesado. Exitosos: {success_count}. Errores: {len(errors)}")
                            if errors:
                                with st.expander("Ver Errores"):
                                    for err in errors:
                                        st.error(err)
                                        
                    except Exception as e:
                        st.error(f"Error al leer archivo: {e}")
            
            with upload_tab2:
                st.subheader("Carga Masiva de Clientes")
                st.info(" **Formato requerido**: Nombre | RUC | Direccin")
                
                template_clients = pd.DataFrame({
                    'Nombre': ['Empresa ABC S.A.C.', 'Corporacin XYZ'],
                    'RUC': ['20123456789', '20987654321'],
                    'Direccin': ['Av. Principal 123, Lima', 'Jr. Secundario 456, Lima']
                })
                # DOWNLOAD TEMPLATES
                if HAS_OPENPYXL:
                    buffer_clients = io.BytesIO()
                    with pd.ExcelWriter(buffer_clients, engine='openpyxl') as writer:
                        template_clients.to_excel(writer, index=False, sheet_name='Clientes')
                    st.download_button(" Descargar Template Clientes", data=buffer_clients.getvalue(), file_name="template_clientes.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                else:
                    st.warning(" Requiere 'openpyxl' para descargar templates.")
                
                uploaded_clients = st.file_uploader("Seleccionar archivo Excel", type=['xlsx'], key="upload_clients")
                if uploaded_clients and HAS_OPENPYXL:
                    try:
                        df_clients = pd.read_excel(uploaded_clients)
                        st.write("Vista previa:", df_clients.head())
                        
                        if st.button("Procesar Carga de Clientes"):
                            success_count = 0
                            for idx, row in df_clients.iterrows():
                                try:
                                    supabase.table("clients").insert({
                                        "name": row.get('Nombre'),
                                        "doi_type": "RUC",
                                        "doi_number": str(row.get('RUC')),
                                        "address": row.get('Direccin', '')
                                    }).execute()
                                    success_count += 1
                                except Exception as e:
                                    st.error(f"Fila {idx+2}: {str(e)}")
                            st.success(f" {success_count} clientes cargados")
                    except Exception as e:
                        st.error(f"Error: {e}")
            
            with upload_tab3:
                st.subheader("Carga Masiva de Proyectos")
                st.info(" **Formato requerido**: Cliente | Nombre Proyecto | Moneda")
                
                template_projects = pd.DataFrame({
                    'Cliente': ['Empresa ABC S.A.C.', 'Corporacin XYZ'],
                    'Nombre Proyecto': ['Implementacin ERP', 'Consultora Fiscal'],
                    'Moneda': ['PEN', 'USD']
                })
                if HAS_OPENPYXL:
                    buffer_projects = io.BytesIO()
                    with pd.ExcelWriter(buffer_projects, engine='openpyxl') as writer:
                        template_projects.to_excel(writer, index=False, sheet_name='Proyectos')
                    st.download_button(" Descargar Template Proyectos", data=buffer_projects.getvalue(), file_name="template_proyectos.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                else:
                    st.warning(" Requiere 'openpyxl' para descargar templates.")
                
                uploaded_projects = st.file_uploader("Seleccionar archivo Excel", type=['xlsx'], key="upload_projects")
                if uploaded_projects and HAS_OPENPYXL:
                    try:
                        df_projects = pd.read_excel(uploaded_projects)
                        st.write("Vista previa:", df_projects.head())
                        
                        if st.button("Procesar Carga de Proyectos"):
                            clients_map = {c['name']: c['id'] for c in supabase.table("clients").select("id, name").execute().data}
                            success_count = 0
                            for idx, row in df_projects.iterrows():
                                try:
                                    cliente = row.get('Cliente')
                                    if cliente not in clients_map:
                                        st.error(f"Fila {idx+2}: Cliente '{cliente}' no encontrado")
                                        continue
                                    
                                    supabase.table("projects").insert({
                                        "client_id": clients_map[cliente],
                                        "name": row.get('Nombre Proyecto'),
                                        "currency": row.get('Moneda', 'PEN')
                                    }).execute()
                                    success_count += 1
                                except Exception as e:
                                    st.error(f"Fila {idx+2}: {str(e)}")
                            st.success(f" {success_count} proyectos cargados")
                    except Exception as e:
                        st.error(f"Error: {e}")
            
            with upload_tab4:
                st.subheader("Carga Masiva de Tarifas")
                st.info(" **Formato requerido**: Proyecto | Rol | Tarifa")
                
                template_rates = pd.DataFrame({
                    'Proyecto': ['Implementacin ERP', 'Consultora Fiscal'],
                    'Rol': ['Consultor Senior', 'Analista'],
                    'Tarifa': [150.00, 80.00]
                })
                if HAS_OPENPYXL:
                    buffer_rates = io.BytesIO()
                    with pd.ExcelWriter(buffer_rates, engine='openpyxl') as writer:
                        template_rates.to_excel(writer, index=False, sheet_name='Tarifas')
                    st.download_button(" Descargar Template Tarifas", data=buffer_rates.getvalue(), file_name="template_tarifas.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                else:
                    st.warning(" Requiere 'openpyxl' para descargar templates.")
                
                uploaded_rates = st.file_uploader("Seleccionar archivo Excel", type=['xlsx'], key="upload_rates")
                if uploaded_rates and HAS_OPENPYXL:
                    try:
                        df_rates = pd.read_excel(uploaded_rates)
                        st.write("Vista previa:", df_rates.head())
                        
                        if st.button("Procesar Carga de Tarifas"):
                            projects_map = {p['name']: p['id'] for p in supabase.table("projects").select("id, name").execute().data}
                            roles_map = {r['name']: r['id'] for r in supabase.table("roles").select("id, name").execute().data}
                            success_count = 0
                            for idx, row in df_rates.iterrows():
                                try:
                                    proyecto = row.get('Proyecto')
                                    rol = row.get('Rol')
                                    
                                    if proyecto not in projects_map:
                                        st.error(f"Fila {idx+2}: Proyecto '{proyecto}' no encontrado")
                                        continue
                                    if rol not in roles_map:
                                        st.error(f"Fila {idx+2}: Rol '{rol}' no encontrado")
                                        continue
                                    
                                    supabase.table("project_rates").insert({
                                        "project_id": projects_map[proyecto],
                                        "role_id": roles_map[rol],
                                        "rate": float(row.get('Tarifa'))
                                    }).execute()
                                    success_count += 1
                                except Exception as e:
                                    st.error(f"Fila {idx+2}: {str(e)}")
                            st.success(f" {success_count} tarifas cargadas")
                    except Exception as e:
                        st.error(f"Error: {e}")

        elif choice == "Facturación y Reportes":
            st.header(" Facturación y Reportes")
            
            # Filtros de Reporte
            clientes_q = supabase.table("clients").select("id, name, doi_type, doi_number, address").order("name").execute()
            if not clientes_q.data:
                st.warning("Debe registrar clientes primero.")
            else:
                with st.sidebar:
                    st.markdown("### Configuracin de Reporte")
                    cli_map = {c['name']: c for c in clientes_q.data}
                    cli_name_sel = st.selectbox("Seleccionar Cliente", list(cli_map.keys()))
                    cli_data = cli_map[cli_name_sel]
                    
                    date_range = st.date_input("Rango de Fechas", [get_lima_now().replace(day=1), get_lima_now()])
                    
                    st.markdown("---")
                    st.info(" Genere el reporte para habilitar la carta y anexos.")

                if len(date_range) == 2:
                    start_d, end_d = date_range
                    report_q = supabase.table("time_entries").select("*, profiles(full_name, role_id, roles(name)), projects(name, currency, client_id)").eq("projects.client_id", cli_data['id']).execute()
                    
                    if report_q.data:
                        df_rep = pd.json_normalize(report_q.data)
                        if df_rep.empty:
                            st.info("No hay registros en el periodo seleccionado.")
                        else:
                            # Procesamiento
                            df_rep['dt_ref'] = df_rep['start_time'].fillna(df_rep['created_at'])
                            df_rep['dt_start'] = (pd.to_datetime(df_rep['dt_ref'], utc=True, errors='coerce') - pd.Timedelta(hours=5)).dt.tz_localize(None)
                            df_rep['Fecha_dt'] = df_rep['dt_start'].dt.date
                            df_rep['Fecha_str'] = df_rep['dt_start'].dt.strftime('%d.%m-%Y').fillna('---')
                            
                            rates = supabase.table("project_rates").select("*").execute()
                            rates_df = pd.DataFrame(rates.data)
                            
                            def get_cost_rep(row):
                                if rates_df.empty: return 0.0
                                r = rates_df[(rates_df['project_id'] == row['project_id']) & (rates_df['role_id'] == row['profiles.role_id'])]
                                return float(r['rate'].iloc[0]) if not r.empty else 0.0
                            
                            df_rep['Horas_num'] = df_rep['total_minutes'] / 60
                            df_rep['Costo_H'] = df_rep.apply(get_cost_rep, axis=1)
                            df_rep['Total_Monto'] = df_rep['Horas_num'] * df_rep['Costo_H']
                            
                            # SELECTOR DE PROYECTOS (Nuevo)
                            st.markdown("###  Selección de Proyectos a Liquidar")
                            proyectos_disponibles = df_rep['projects.name'].unique().tolist()
                            proyectos_seleccionados = st.multiselect(
                                "Seleccione los proyectos que desea incluir en esta liquidación:",
                                options=proyectos_disponibles,
                                default=proyectos_disponibles  # Por defecto todos seleccionados
                            )
                            
                            if proyectos_seleccionados:
                                # Filtrar dataframe por proyectos seleccionados
                                df_rep = df_rep[df_rep['projects.name'].isin(proyectos_seleccionados)]
                            
                                tab1, tab2, tab3 = st.tabs([" Carta de Liquidación", " Anexo Detallado", " Dashboard"])
                            
                                with tab1:
                                    monedas_disp = [m for m in df_rep['projects.currency'].unique() if pd.notna(m) and str(m) != 'nan']
                                    if not monedas_disp:
                                        st.warning("No hay monedas vlidas.")
                                    else:
                                        moneda_liq = st.selectbox("Moneda para Carta", monedas_disp)
                                        df_carta = df_rep[df_rep['projects.currency'] == moneda_liq].copy()
                                        total_general_liq = df_carta['Total_Monto'].sum()
                                        
                                        # Datos Pre-llenados
                                        doi_str = str(cli_data.get('doi_number', '')).strip()
                                        if doi_str == 'nan' or not doi_str: doi_str = '---'
                                        addr_str = str(cli_data.get('address', '')).strip()
                                        if addr_str == 'nan' or not addr_str: addr_str = 'Lima, Per.'
                                        try:
                                            firma_def = st.session_state.profile['full_name']
                                        except: firma_def = "Responsable"
                                        
                                        # ===== CONTROL DE LIQUIDACIN =====
                                        st.markdown("---")
                                        st.markdown("###  Control de Liquidación")
                                        
                                        # Verificar liquidación existente
                                        existing_liq = supabase.table("liquidations").select("*").eq("client_id", cli_data['id']).eq("period_start", start_d.isoformat()).eq("period_end", end_d.isoformat()).eq("currency", moneda_liq).execute()
                                        
                                        liquidation_number = None
                                        liquidation_id = None
                                        liquidation_status = "draft"
                                        
                                        if existing_liq.data:
                                            liq_data = existing_liq.data[0]
                                            liquidation_number = liq_data['liquidation_number']
                                            liquidation_id = liq_data['id']
                                            liquidation_status = liq_data.get('status', 'draft')
                                            st.info(f" Liquidación existente: **{liquidation_number}** | Estado: **{liquidation_status.upper()}**")
                                        else:
                                            st.caption(" No se ha generado número de liquidación. Se generará al guardar.")
                                        
                                        # Campo para notas especiales (descuentos, condiciones, etc.)
                                        st.markdown("#####  Notas Especiales (Opcional)")
                                        st.caption("Agregue aqu descuentos, condiciones especiales o cualquier texto adicional que desee incluir en la carta.")
                                        notas_especiales = st.text_area(
                                            "Notas adicionales para esta liquidación:",
                                            placeholder="Ejemplo: Se aplic un descuento del 10% por volumen de horas.\nO: Monto neto a pagar: USD 5,400.00 (despus de descuento de USD 600.00)",
                                            height=100,
                                            key=f"notas_{cli_name_sel}_{moneda_liq}",
                                            value=existing_liq.data[0].get('special_notes', '') if existing_liq.data else ''
                                        )
                                        
                                        # Construir seccin de notas si existe
                                        seccion_notas = ""
                                        if notas_especiales and notas_especiales.strip():
                                            seccion_notas = f"\n\n{notas_especiales.strip()}"
                                        
                                        # Construir referencia con número
                                        ref_line = "Ref.: Liquidación de Honorarios"
                                        if liquidation_number:
                                            ref_line = f"Ref.: Liquidación de Honorarios N {liquidation_number}"
                                        
                                        # Plantilla de Carta basada en PDF Hoja 1
                                        fecha_carta = get_lima_now().strftime('%d de %B de %Y')
                                        
                                        letter_template = f"""San Isidro, {fecha_carta}

Seor(es):
{cli_name_sel.upper()}
Presente.-

Estimado(s) seor(es):

Nos dirigimos a usted(es) con el propsito de saludarlo(s) cordialmente y remitir la {ref_line}, por la suma neta de {moneda_liq} {total_general_liq:,.2f}, ms el Impuesto General a las Ventas.

El detalle de las actividades efectivamente ejecutadas a favor de usted(es) se encuentra consignado en la liquidación de horas que se adjunta a esta comunicación. En tal sentido, agradeceremos se sirvan revisar detenidamente la información anexada.{seccion_notas}

Para el pago de los honorarios y de la respectiva detraccin, srvanse tener en cuenta los siguientes datos:

**Pago Detracciones:** Banco de la Nacin
Cuenta corriente Soles N 00-005-337240

**Pago Honorarios:** 
[BANCO] [TIPO CUENTA] 
N [NUMERO DE CUENTA]

Atentamente,

__________________________
{firma_def}
Responsable"""
                                        
                                        st.markdown("##### Editor de Carta")
                                        full_letter_text = st.text_area("Contenido", value=letter_template, height=450)
                                        
                                        # Botones para guardar liquidación
                                        st.markdown("---")
                                        col_save1, col_save2, col_save3 = st.columns([1, 1, 1])
                                        
                                        with col_save1:
                                            if st.button(" Guardar Liquidación", type="primary", help="Guardar liquidación y generará número correlativo"):
                                                try:
                                                    # Generar número si no existe
                                                    if not liquidation_number:
                                                        result = supabase.rpc('get_next_liquidation_number').execute()
                                                        liquidation_number = result.data
                                                    
                                                    # Preparar datos
                                                    liq_data_to_save = {
                                                        "client_id": cli_data['id'],
                                                        "period_start": start_d.isoformat(),
                                                        "period_end": end_d.isoformat(),
                                                        "currency": moneda_liq,
                                                        "total_amount": float(total_general_liq),
                                                        "special_notes": notas_especiales.strip() if notas_especiales else None,
                                                        "projects": proyectos_seleccionados,
                                                        "status": "draft",
                                                        "generated_by": st.session_state.user.id
                                                    }
                                                    
                                                    if liquidation_id:
                                                        supabase.table("liquidations").update(liq_data_to_save).eq("id", liquidation_id).execute()
                                                        st.success(f" Liquidación {liquidation_number} actualizada")
                                                    else:
                                                        liq_data_to_save["liquidation_number"] = liquidation_number
                                                        supabase.table("liquidations").insert(liq_data_to_save).execute()
                                                        st.success(f" Liquidación {liquidation_number} guardada")
                                                    
                                                    st.rerun()
                                                except Exception as e:
                                                    st.error(f"Error al guardar: {str(e)}")
                                        
                                        with col_save2:
                                            if liquidation_id and liquidation_status == "draft":
                                                if st.button(" Marcar como Enviada"):
                                                    try:
                                                        supabase.table("liquidations").update({"status": "sent", "sent_at": get_lima_now().astimezone(timezone.utc).replace(tzinfo=None).isoformat()}).eq("id", liquidation_id).execute()
                                                        st.success(" Marcada como Enviada")
                                                        st.rerun()
                                                    except Exception as e:
                                                        st.error(f"Error: {str(e)}")
                                        
                                        with col_save3:
                                            if liquidation_id and liquidation_status == "sent":
                                                if st.button(" Marcar como Pagada"):
                                                    try:
                                                        supabase.table("liquidations").update({"status": "paid", "paid_at": get_lima_now().astimezone(timezone.utc).replace(tzinfo=None).isoformat()}).eq("id", liquidation_id).execute()
                                                        st.success(" Marcada como Pagada")
                                                        st.rerun()
                                                    except Exception as e:
                                                        st.error(f"Error: {str(e)}")
                                        
                                        st.markdown("---")
                                        c1, c2 = st.columns(2)
                                        with c1:
                                            st.caption("Vista Previa")
                                            st.markdown(f"<div style='background:white; color:black; padding:25px; border:1px solid #ccc; font-family:Times New Roman; white-space: pre-wrap;'>{full_letter_text}</div>", unsafe_allow_html=True)
                                        with c2:
                                            st.caption("Acciones")
                                            if HAS_DOCX:
                                                docx_bytes = generate_word_letter(full_letter_text, firma_def)
                                                st.download_button(" Descargar Word (.docx)", data=docx_bytes, file_name=f"Carta_{cli_name_sel}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
                                            else:
                                                st.warning("Instale python-docx.")

                                with tab2:
                                    if 'moneda_liq' in locals() and moneda_liq:
                                        st.subheader(f"Anexo: Detalle ({moneda_liq})")
                                        df_anexo = df_rep[df_rep['projects.currency'] == moneda_liq].copy()
                                        full_xls = []
                                        for proj in df_anexo['projects.name'].unique():
                                            st.markdown(f"**Proyecto: {proj}**")
                                            df_p = df_anexo[df_anexo['projects.name'] == proj].copy()
                                            disp = df_p[['Fecha_str', 'profiles.full_name', 'description', 'total_minutes', 'Total_Monto']].copy()
                                            disp['Tiempo'] = disp['total_minutes'].apply(lambda x: f"{int(x)//60:02d}:{int(x)%60:02d}")
                                            disp = disp[['Fecha_str', 'profiles.full_name', 'description', 'Tiempo', 'Total_Monto']]
                                            disp.columns = ['Fecha', 'Consultor', 'Actividad', 'Tiempo', 'Valor']
                                            st.dataframe(disp, column_config={"Valor": st.column_config.NumberColumn(format="%.2f")}, use_container_width=True, hide_index=True)
                                            
                                            clean_p = disp.copy()
                                            clean_p.insert(0, 'Proyecto', proj)
                                            full_xls.append(clean_p)
                                        
                                        st.markdown("---")
                                        if HAS_OPENPYXL and full_xls:
                                            try:
                                                final_xls = pd.concat(full_xls)
                                                buffer = io.BytesIO()
                                                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                                                    final_xls.to_excel(writer, index=False, sheet_name='Anexo')
                                                st.download_button(f" Descargar Anexo Detallado ({moneda_liq})", data=buffer.getvalue(), file_name=f"anexo_{cli_name_sel}_{moneda_liq}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                                            except Exception as e:
                                                st.error(f"Error generando Excel: {str(e)}")
                                        elif full_xls:
                                            st.warning(" Requiere 'openpyxl' para descargar el anexo en Excel.")
                                    else:
                                        st.info("Seleccione moneda en pestaa Carta.")

                                with tab3:
                                    st.subheader("Dashboard")
                                    sum_df = df_rep.groupby(['profiles.full_name', 'projects.currency'])[['Horas_num', 'Total_Monto']].sum().reset_index()
                                    sum_df['Tiempo'] = sum_df['Horas_num'].apply(lambda h: f"{int(h)}h {int((h*60)%60)}m")
                                    st.dataframe(sum_df, column_config={"Total_Monto": st.column_config.NumberColumn(format="%.2f")}, use_container_width=True, hide_index=True)
                            else:
                                st.warning(" Debe seleccionar al menos un proyecto para generará la liquidación.")

                    else:
                        st.info("No se encontraron registros para este cliente.")
                else:
                    st.info("Seleccione un rango de fechas en la barra lateral.")

    else:
        # Para roles de usuario no administrador
        mostrar_registro_tiempos()

# --- REFRESH DINMICO (Al final para no bloquear UI) ---

if st.session_state.get('user') and st.session_state.get('timer_running') and not st.session_state.get('logout_requested'):
    time.sleep(1)
    st.rerun()
